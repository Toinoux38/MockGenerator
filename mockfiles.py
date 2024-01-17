"""
Author : Antoine.H
fakeGenerator : Generate fakes mock files | DOCX, PDF and XLSX | Custom logo 
"""

import os
import random
from faker import Faker
from reportlab.pdfgen import canvas
from docx import Document
from openpyxl import Workbook
from datetime import datetime, timedelta
from docx.shared import Pt
from docx.oxml import OxmlElement
from rich.console import Console
from rich.progress import Progress


ascii_art = '''

   __      _         _____                           _             
  / _|    | |       / ____|                         | |            
 | |_ __ _| | _____| |  __  ___ _ __   ___ _ __ __ _| |_ ___  _ __ 
 |  _/ _` | |/ / _ \ | |_ |/ _ \ '_ \ / _ \ '__/ _` | __/ _ \| '__|
 | || (_| |   <  __/ |__| |  __/ | | |  __/ | | (_| | || (_) | |   
 |_| \__,_|_|\_\___|\_____|\___|_| |_|\___|_|  \__,_|\__\___/|_|   
                                                                   
                                                                                                          
'''

print(ascii_art)
print("made with love by macron")
faker = Faker()

def generate_pdf(file_name):
    with open(file_name, 'wb') as file:
        pdf = canvas.Canvas(file_name)

        pdf.drawString(100, 750, f"Name: {faker.name()}")
        pdf.drawString(100, 730, f"Address: {faker.address()}")
        pdf.drawString(100, 710, f"Email: {faker.email()}")


        salary = random.randint(30000, 80000)
        bonus = random.randint(1000, 5000)
        expenses = random.randint(500, 2000)
        total_income = salary + bonus

        table_data = [
            ['Category', 'Amount'],
            ['Salary', f"${salary:,}"],
            ['Bonus', f"${bonus:,}"],
            ['Expenses', f"${expenses:,}"],
            ['Total Income', f"${total_income:,}"],
        ]

        col_widths = [pdf.stringWidth(str(max(data, key=len)), "Helvetica", 12) + 6 for data in zip(*table_data)]

        x = 100
        y = 600

        for row in table_data:
            for i, cell in enumerate(row):
                pdf.drawString(x, y, cell)
                x += col_widths[i]
            x = 100
            y -= 20

        # CHANGE LOGO HERE ############
        logo_path = 'burger_king_logo.png' 
        pdf.drawImage(logo_path, x=100, y=100, width=70, height=70)

        # Add the name "Burger King"
        pdf.setFont("Helvetica", 12)
        pdf.drawString(100, 80, "Burger King")

        pdf.save()

def generate_docx(file_name):
    document = Document()

    title = document.add_heading("Gestion des stocks", level=1)
    title.paragraph_format.alignment = 1 

    logo_path = 'burger_king_logo.jpg'  # CHANGE LOGO HERE ##############
    paragraph = document.add_paragraph()
    run = paragraph.add_run()
    run.add_picture(logo_path, width=Pt(100), height=Pt(100))
    paragraph.alignment = 1 

    table = document.add_table(rows=1, cols=5)
    for row in table.rows:
        for cell in row.cells:
            cell.paragraphs[0].paragraph_format.alignment = 1 


    header_cells = table.rows[0].cells
    header_cells[0].text = 'Product'
    header_cells[1].text = 'Quantity'
    header_cells[2].text = 'Price'
    header_cells[3].text = 'Supplier'
    header_cells[4].text = 'Last Updated'

    for _ in range(10):
        product = faker.word()
        quantity = random.randint(1, 100)
        price = round(random.uniform(5.0, 50.0), 2)
        supplier = faker.company()
        last_updated = faker.date_this_year()

        row_cells = table.add_row().cells
        row_cells[0].text = product
        row_cells[1].text = str(quantity)
        row_cells[2].text = f"${price:,}"
        row_cells[3].text = supplier
        row_cells[4].text = last_updated.strftime('%Y-%m-%d')

    document.save(file_name)


def generate_xlsx(file_name):
    wb = Workbook()
    ws = wb.active

    current_date = datetime.now().strftime('%Y-%m-%d')
    ws.append([f"Date: {current_date}"])
    headers = ['Client Name', 'Order Number', 'Order Content', 'Price', 'Delivered', 'Employee']
    ws.append(headers)

    for _ in range(400):
        client_name = faker.name()
        order_number = random.randint(100000000000000, 999999999999999)  # Generate a UUID (yes this is a shit way)
        order_content = faker.sentence()
        price = round(random.uniform(5.0, 50.0), 2) 
        delivered = random.choice(['Yes', 'No'])
        employee = faker.first_name() 

        row_data = [client_name, order_number, order_content, price, delivered, employee]
        ws.append(row_data)
        
    wb.save(file_name)

# Create directories
pdf_dir = 'pdf_files'
docx_dir = 'docx_files'
xlsx_dir = 'xlsx_files'

os.makedirs(pdf_dir, exist_ok=True)
os.makedirs(docx_dir, exist_ok=True)
os.makedirs(xlsx_dir, exist_ok=True)


total_files = 100
with Progress() as progress:
    task = progress.add_task("[cyan]Generating files...", total=total_files)
    for i in range(1, total_files + 1):
        progress.update(task, advance=1)
        pdf_file = os.path.join(pdf_dir, f'mock_data_{i}.pdf')
        docx_file = os.path.join(docx_dir, f'gestion_stocks_{i}.docx')
        xlsx_file = os.path.join(xlsx_dir, f'commandes_bk{i}.xlsx')

        generate_pdf(pdf_file)
        generate_docx(docx_file)
        generate_xlsx(xlsx_file)
    progress.stop()

console = Console()
console.print("[bold green]Mock files generated successfully.")
